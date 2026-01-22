"""
Document parsers for different file types.
Extracts text while preserving structure (page numbers, sections, etc.).
"""
import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
from typing import Dict, List, Tuple
import re


class PDFParser:
    """Parse PDF files and extract text with page numbers."""
    
    @staticmethod
    def parse(file_path: str) -> List[Dict]:
        """
        Parse PDF and return list of page dictionaries.
        
        Returns:
            List of dicts with 'page', 'text', 'metadata'
        """
        doc = fitz.open(file_path)
        pages = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            # Use "text" to get plain text, and replace common odd characters
            text = page.get_text("text").encode("ascii", "ignore").decode("ascii")
            
            pages.append({
                'page': page_num + 1,
                'text': text.strip(),
                'metadata': {
                    'page': page_num + 1,
                    'total_pages': len(doc)
                }
            })
        
        doc.close()
        return pages


class DOCXParser:
    """Parse DOCX files and extract text with section headings."""
    
    @staticmethod
    def parse(file_path: str) -> List[Dict]:
        """
        Parse DOCX and return list of section dictionaries.
        
        Returns:
            List of dicts with 'section', 'text', 'metadata'
        """
        doc = Document(file_path)
        sections = []
        current_section = "Introduction"
        current_text = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Detect headings (simple heuristic: short lines, possibly numbered)
            if len(text) < 100 and (para.style.name.startswith('Heading') or 
                                    re.match(r'^(\d+\.|\d+\)|\w+\.)\s+[A-Z]', text)):
                # Save previous section
                if current_text:
                    sections.append({
                        'section': current_section,
                        'text': '\n'.join(current_text),
                        'metadata': {
                            'section': current_section
                        }
                    })
                    current_text = []
                
                current_section = text
            else:
                current_text.append(text)
        
        # Add final section
        if current_text:
            sections.append({
                'section': current_section,
                'text': '\n'.join(current_text),
                'metadata': {
                    'section': current_section
                }
            })
        
        return sections if sections else [{'section': 'Document', 'text': '\n'.join([p.text for p in doc.paragraphs]), 'metadata': {}}]


class TextParser:
    """Parse plain text and markdown files."""
    
    @staticmethod
    def parse(file_path: str) -> List[Dict]:
        """
        Parse text/markdown files.
        
        Returns:
            List of dicts with 'section', 'text', 'metadata'
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Try to detect markdown sections
        if file_path.endswith('.md'):
            return TextParser._parse_markdown(content)
        else:
            return [{'section': 'Document', 'text': content, 'metadata': {}}]
    
    @staticmethod
    def _parse_markdown(content: str) -> List[Dict]:
        """Parse markdown with section detection."""
        sections = []
        lines = content.split('\n')
        current_section = "Introduction"
        current_text = []
        
        for line in lines:
            # Detect markdown headings
            if line.startswith('#'):
                # Save previous section
                if current_text:
                    sections.append({
                        'section': current_section,
                        'text': '\n'.join(current_text),
                        'metadata': {'section': current_section}
                    })
                    current_text = []
                
                current_section = line.lstrip('#').strip()
            else:
                current_text.append(line)
        
        # Add final section
        if current_text:
            sections.append({
                'section': current_section,
                'text': '\n'.join(current_text),
                'metadata': {'section': current_section}
            })
        
        return sections if sections else [{'section': 'Document', 'text': content, 'metadata': {}}]


class DocumentParser:
    """Main parser that routes to appropriate parser based on file type."""
    
    @staticmethod
    def parse(file_path: str) -> List[Dict]:
        """
        Parse document and return structured content.
        
        Args:
            file_path: Path to the document
            
        Returns:
            List of content dictionaries with text and metadata
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension == '.pdf':
            return PDFParser.parse(file_path)
        elif extension == '.docx':
            return DOCXParser.parse(file_path)
        elif extension in ['.txt', '.md']:
            return TextParser.parse(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
